"""DOCX Exporter — Converts markdown proposals to professional Word documents.

Uses python-docx to generate .docx files with:
- Cover page with company branding
- Table of Contents placeholder
- Styled headings, paragraphs, and tables
- Headers and footers with page numbers
- Bullet and numbered lists
"""

import re
import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT


# ── Brand Colors ───────────────────────────────────────────────────

BRAND_PRIMARY = RGBColor(0x1A, 0x56, 0xDB)     # Blue
BRAND_SECONDARY = RGBColor(0x2D, 0x3A, 0x4A)   # Dark gray-blue
BRAND_ACCENT = RGBColor(0x10, 0xB9, 0x81)       # Green
BRAND_LIGHT_BG = RGBColor(0xF0, 0xF4, 0xF8)     # Light background
BRAND_TEXT = RGBColor(0x33, 0x33, 0x33)          # Body text


# ── Document Setup ─────────────────────────────────────────────────

def _setup_styles(doc: Document):
    """Configure document styles for professional look."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = BRAND_TEXT
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = BRAND_PRIMARY
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = BRAND_SECONDARY
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = BRAND_SECONDARY
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)


def _add_cover_page(doc: Document, client_name: str, project_title: str, company_name: str = "AZA FUTURE"):
    """Add a professional cover page."""
    # Spacing before title
    for _ in range(6):
        doc.add_paragraph()

    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company_name.upper())
    run.font.size = Pt(14)
    run.font.color.rgb = BRAND_PRIMARY
    run.font.bold = True
    run.font.name = "Calibri"

    # Divider line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 40)
    run.font.color.rgb = BRAND_PRIMARY
    run.font.size = Pt(12)

    # Project title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(project_title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = BRAND_SECONDARY
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(12)

    # Client name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Prepared for: {client_name}")
    run.font.size = Pt(16)
    run.font.color.rgb = BRAND_TEXT
    run.font.name = "Calibri"

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datetime.now().strftime("%B %d, %Y"))
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.name = "Calibri"

    # Confidential notice
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONFIDENTIAL")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    run.font.name = "Calibri"
    run.font.italic = True

    # Page break after cover
    doc.add_page_break()


def _add_header_footer(doc: Document, company_name: str = "AZA FUTURE", client_name: str = ""):
    """Add headers and footers to all sections."""
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"{company_name} — Proposal for {client_name}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        run.font.name = "Calibri"

        # Footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Confidential — ")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        run.font.name = "Calibri"

        # Page number field
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run2 = p.add_run()
        run2._element.append(fldChar1)

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run3 = p.add_run()
        run3.font.size = Pt(8)
        run3.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        run3._element.append(instrText)

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run4 = p.add_run()
        run4._element.append(fldChar2)


# ── Markdown Parser ────────────────────────────────────────────────

def _parse_markdown_to_docx(doc: Document, markdown_text: str):
    """Convert markdown content to docx elements."""
    lines = markdown_text.split("\n")
    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Skip horizontal rules and empty decorative lines
        if re.match(r"^---+$", line.strip()) or re.match(r"^\*\*\*+$", line.strip()):
            i += 1
            continue

        # Skip metadata lines (Generated: ... | Language: ...)
        if line.strip().startswith("*Generated:") or line.strip().startswith("_Generated:"):
            i += 1
            continue

        # Table detection
        if "|" in line and not in_table:
            # Check if this is a table (has separator row)
            if i + 1 < len(lines) and re.match(r"^\|[-\s|:]+\|$", lines[i + 1].strip()):
                in_table = True
                table_rows = [line]
                i += 1
                continue

        if in_table:
            if "|" in line:
                # Skip separator row
                if not re.match(r"^\|[-\s|:]+\|$", line.strip()):
                    table_rows.append(line)
                i += 1
                continue
            else:
                # End of table — render it
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
                # Don't increment, process current line normally

        # Headings
        if line.startswith("# "):
            text = _clean_markdown(line[2:].strip())
            doc.add_heading(text, level=1)
            i += 1
            continue

        if line.startswith("## "):
            text = _clean_markdown(line[3:].strip())
            doc.add_heading(text, level=2)
            i += 1
            continue

        if line.startswith("### "):
            text = _clean_markdown(line[4:].strip())
            doc.add_heading(text, level=3)
            i += 1
            continue

        # Bullet lists
        if re.match(r"^[-*]\s+", line):
            text = _clean_markdown(re.sub(r"^[-*]\s+", "", line))
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_text(p, text)
            i += 1
            continue

        # Numbered lists
        if re.match(r"^\d+\.\s+", line):
            text = _clean_markdown(re.sub(r"^\d+\.\s+", "", line))
            p = doc.add_paragraph(style="List Number")
            _add_rich_text(p, text)
            i += 1
            continue

        # Empty lines
        if not line.strip():
            i += 1
            continue

        # Regular paragraphs
        text = _clean_markdown(line.strip())
        if text:
            p = doc.add_paragraph()
            _add_rich_text(p, text)

        i += 1

    # Flush remaining table
    if in_table and table_rows:
        _add_table(doc, table_rows)


def _clean_markdown(text: str) -> str:
    """Remove markdown decorators but keep text structure."""
    # Remove emoji at start of headings
    text = re.sub(r"^[📄📋📅💰📌⚠️🚀✅📊📝🔍📂]+\s*", "", text)
    return text.strip()


def _add_rich_text(paragraph, text: str):
    """Parse inline markdown (bold, italic) into Word runs."""
    # Split by bold markers
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Check for italic
            sub_parts = re.split(r"(\*[^*]+\*)", part)
            for sub in sub_parts:
                if sub.startswith("*") and sub.endswith("*") and len(sub) > 2:
                    run = paragraph.add_run(sub[1:-1])
                    run.italic = True
                else:
                    if sub:
                        paragraph.add_run(sub)


def _add_table(doc: Document, rows: list[str]):
    """Convert markdown table rows into a Word table."""
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)

    if not parsed:
        return

    num_cols = len(parsed[0])
    table = doc.add_table(rows=len(parsed), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    for row_idx, row_data in enumerate(parsed):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.cell(row_idx, col_idx)
                cell.text = ""
                p = cell.paragraphs[0]
                _add_rich_text(p, _clean_markdown(cell_text))

                # Bold header row
                if row_idx == 0:
                    for run in p.runs:
                        run.bold = True

    doc.add_paragraph()  # Space after table


# ── Public API ─────────────────────────────────────────────────────

def export_proposal_to_docx(
    markdown_content: str,
    client_name: str,
    project_title: str = "Technical Proposal",
    output_path: str | None = None,
    company_name: str = "AZA FUTURE",
) -> str:
    """Export a markdown proposal to a professional DOCX file.

    Args:
        markdown_content: The proposal in markdown format
        client_name: Name of the client company
        project_title: Title for the cover page
        output_path: Where to save (auto-generated if None)
        company_name: Your company name for branding

    Returns:
        str: Path to the generated .docx file
    """
    doc = Document()

    # Page setup (A4)
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Setup
    _setup_styles(doc)
    _add_cover_page(doc, client_name, project_title, company_name)
    _add_header_footer(doc, company_name, client_name)

    # Parse and add content
    _parse_markdown_to_docx(doc, markdown_content)

    # Save
    if output_path is None:
        safe_name = re.sub(r"[^a-zA-Z0-9]", "_", client_name.lower())
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = Path("exports")
        output_dir.mkdir(exist_ok=True)
        output_path = str(output_dir / f"proposal_{safe_name}_{timestamp}.docx")

    doc.save(output_path)
    return output_path